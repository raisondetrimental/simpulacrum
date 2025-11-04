import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with formatting"""

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    in_table = False
    table_data = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines unless we're in a table
        if not line and not in_table:
            i += 1
            continue

        # Main title (# )
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = doc.add_heading(title_text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.bold = True

        # Section headers (## )
        elif line.startswith('## ') and not line.startswith('### '):
            header_text = line[3:].strip()
            p = doc.add_heading(header_text, level=1)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 51, 102)

        # Subsection headers (### )
        elif line.startswith('### ') and not line.startswith('#### '):
            header_text = line[4:].strip()
            p = doc.add_heading(header_text, level=2)
            for run in p.runs:
                run.font.size = Pt(14)

        # Sub-subsection headers (#### )
        elif line.startswith('#### '):
            header_text = line[5:].strip()
            p = doc.add_heading(header_text, level=3)
            for run in p.runs:
                run.font.size = Pt(12)

        # Horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        # Tables
        elif line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_data = []

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)

            # Check if next line is separator or end of table
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if not next_line.startswith('|'):
                    # End of table - create it
                    if len(table_data) > 1:  # At least header + separator
                        # Remove separator row (usually second row with dashes)
                        if all('---' in cell or ':-' in cell or '-:' in cell for cell in table_data[1]):
                            table_data.pop(1)

                        # Create table
                        if table_data:
                            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            table.style = 'Light Grid Accent 1'

                            # Populate table
                            for row_idx, row_data in enumerate(table_data):
                                for col_idx, cell_text in enumerate(row_data):
                                    cell = table.rows[row_idx].cells[col_idx]
                                    cell.text = cell_text

                                    # Bold header row
                                    if row_idx == 0:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.font.bold = True

                            doc.add_paragraph()  # Space after table

                    in_table = False
                    table_data = []

        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Handle bold text **text**
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)

        # Regular paragraphs
        elif line.strip():
            text = line.strip()

            # Skip table separators
            if all(c in '|-: ' for c in text):
                i += 1
                continue

            # Handle bold text **text**
            p = doc.add_paragraph()

            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)

            p.paragraph_format.space_after = Pt(6)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Successfully converted to {docx_file}")

if __name__ == "__main__":
    md_file = "Uzbekistan_Market_Intelligence_Report.md"
    docx_file = "Uzbekistan_Market_Intelligence_Report.docx"
    convert_markdown_to_docx(md_file, docx_file)
